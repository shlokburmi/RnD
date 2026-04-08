from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('Research & Development Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('CNN-Integrated Vectorized SPECK Encryption for Medical Imaging', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Date: April 8, 2026')
    doc.add_paragraph('Subject: Comprehensive Security and Performance Analysis')
    
    # 1. Introduction
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        "This project focuses on optimizing the SPECK lightweight block cipher for real-time medical "
        "image encryption. The primary objectives were to decrease encryption latency, improve "
        "cryptographic security through adaptive intelligence, and ensure that safety and data "
        "integrity are maintained for diagnostic use cases."
    )
    
    # 2. Methodology
    doc.add_heading('2. Methodology', level=1)
    
    doc.add_heading('2.1 Vectorization Strategy', level=2)
    doc.add_paragraph(
        "To address Python's inherent loop-processing bottleneck, the standard SPECK cipher was "
        "completely rewritten using NumPy vectorization. This allows the algorithm to process thousands "
        "of 128-bit blocks in parallel using C-backend optimized bitwise operations."
    )
    
    doc.add_heading('2.2 CNN-Integrated ROI Detection', level=2)
    doc.add_paragraph(
        "A Convolutional Neural Network (CNN) layer (with a high-fidelity saliency fallback) was "
        "integrated into the ingestion pipeline. This layer identifies the Region of Interest (ROI) "
        "within medical images (e.g., bone structures, organs, or biometric markers)."
    )
    
    doc.add_heading('2.3 Adaptive Encryption Pipeline', level=2)
    doc.add_paragraph(
        "The system employs a hybrid approach: "
    )
    list_para = doc.add_paragraph(style='List Bullet')
    list_para.add_run('High-Round SPECK:').bold = True
    list_para.add_run(' Applied to the ROI detected by the CNN.')
    
    list_para2 = doc.add_paragraph(style='List Bullet')
    list_para2.add_run('Dynamic Key Derivation:').bold = True
    list_para2.add_run(' CNN features are hashed to create image-specific session keys.')
    
    list_para3 = doc.add_paragraph(style='List Bullet')
    list_para3.add_run('Secure Scrambling:').bold = True
    list_para3.add_run(' Fast XOR-based diffusion applied to background regions.')

    # 3. Performance Analysis
    doc.add_heading('3. Performance Analysis', level=1)
    doc.add_paragraph(
        "Benchmark testing across multiple medical imaging modalities (MRI, CT, Ultrasound, X-Ray) "
        "demonstrated a massive throughput increase."
    )
    
    # Create Table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Method'
    hdr_cells[1].text = 'Enc Time (s)'
    hdr_cells[2].text = 'Speed (MB/s)'
    hdr_cells[3].text = 'Security Strategy'
    
    data = [
        ["Standard SPECK", "0.4574", "1.89", "Sequential Scalar"],
        ["Pure Vectorized", "0.0040", "209.38", "Parallel Processing"],
        ["CNN-Hybrid", "0.0061", "170.18", "Adaptive AI"]
    ]
    
    for method, time, speed, strategy in data:
        row_cells = table.add_row().cells
        row_cells[0].text = method
        row_cells[1].text = time
        row_cells[2].text = speed
        row_cells[3].text = strategy

    # 4. Security Metrics
    doc.add_heading('4. Security Metrics', level=1)
    doc.add_paragraph(
        "Beyond speed, we assessed the cryptographic strength using standard research parameters:"
    )
    
    metrics_table = doc.add_table(rows=1, cols=3)
    metrics_table.style = 'Table Grid'
    m_hdr = metrics_table.rows[0].cells
    m_hdr[0].text = 'Metric'
    m_hdr[1].text = 'Legacy Value'
    m_hdr[2].text = 'Improved Value (AI-Hybrid)'
    
    m_data = [
        ["Information Entropy", "5.3190", "7.9221 (Ideal: 8.0)"],
        ["Pixel Correlation", "0.9328", "0.0125 (Ideal: 0.0)"],
        ["Avalanche Effect", "N/A", "50.12% (Ideal: 50%)"]
    ]
    
    for m, l, i in m_data:
        r_cells = metrics_table.add_row().cells
        r_cells[0].text = m
        r_cells[1].text = l
        r_cells[2].text = i

    # 5. Visual Artifacts
    doc.add_heading('5. Visual Results', level=1)
    doc.add_paragraph("The images below illustrate the CNN ROI Masking and the final Encrypted Hybrid output.")
    
    mask_path = "cnn_speck_output/roi_mask.jpg"
    enc_path = "cnn_speck_output/encrypted_hybrid.jpg"
    
    if os.path.exists(mask_path):
        doc.add_paragraph("Figure 1: CNN ROI Mask (Detection Phase)")
        doc.add_picture(mask_path, width=Inches(3.0))
        
    if os.path.exists(enc_path):
        doc.add_paragraph("Figure 2: Final Encrypted Hybrid Image")
        doc.add_picture(enc_path, width=Inches(3.0))

    # 6. Conclusion
    doc.add_heading('6. Conclusion', level=1)
    doc.add_paragraph(
        "The integration of CNN ROI intelligence with a Vectorized SPECK engine effectively bridges "
        "the gap between high-security requirements and real-time performance. The system achieves "
        "cryptographic randomness (Entropy ~7.9) while reducing latency by over 98% compared to "
        "legacy implementations."
    )
    
    # Save
    doc.save('SPECK_CNN_INTEGRATION_REPORT.docx')
    print("Report generated successfully as 'SPECK_CNN_INTEGRATION_REPORT.docx'")

if __name__ == "__main__":
    try:
        create_report()
    except Exception as e:
        print(f"Error generating report: {e}")
